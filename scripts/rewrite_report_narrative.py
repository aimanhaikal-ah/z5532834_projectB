"""Rewrite the main narrative in report/report.docx for the Part B brief."""
from __future__ import annotations

import pathlib
import re

from docx import Document
from docx.shared import Pt


ROOT = pathlib.Path(__file__).resolve().parent.parent
REPORT = ROOT / "report" / "report.docx"


NARRATIVE = [
    ("Heading 1", "Introduction"),
    (
        "Normal",
        "BetaVest is an evidence-based fund comparison dashboard for young retail investors. "
        "Part A created the clean data foundation: 50 US equities, 10 cryptocurrencies, and "
        "equity news headlines aligned to trading dates over 2020-2023. Part B turns that "
        "foundation into a set of investable fund tests, a standalone sector sentiment index, "
        "a sentiment-fusion portfolio, and a Streamlit app that lets a user inspect the "
        "evidence behind each fund.",
    ),
    (
        "Normal",
        "The report follows the Part B brief by moving from fund construction to investor use. "
        "It first explains the funds and out-of-sample backtest, then compares the fund results "
        "and fact sheets. It next evaluates the sector sentiment index, sentiment fusion, and "
        "robustness tests before examining the Streamlit investor journey and the improvements "
        "that remain.",
    ),
    (
        "Normal",
        "The main finding is that Combined Risk Parity is the most credible core BetaVest fund "
        "in the live backtest. It does not have the highest raw return, but it has the strongest "
        "risk-adjusted result and a more acceptable drawdown profile than the crypto-only funds. "
        "The sentiment extension also adds value, but only as a controlled equity tilt. Equity "
        "The finance-lexicon tilt selected on 2021-2022 improves the Equity Risk Parity Sharpe "
        "ratio from 0.944 to 1.003 in the untouched 2023 holdout. The gain is useful, but it is "
        "not strong enough to replace diversification or risk control.",
    ),
    (
        "Normal",
        "This framing is important for the target user. A young retail investor may be attracted "
        "to high-return crypto strategies, but the evidence shows that return alone can be "
        "misleading. The report and app therefore present funds by investor use case: core, "
        "aggressive, defensive, or speculative. The goal is not to claim that one strategy will "
        "win in the future. The goal is to make the trade-off between return, drawdown, cost, "
        "diversification, and sentiment visible.",
    ),
    (
        "Normal",
        "A second purpose of BetaVest is to make technical fund evidence usable for a non-technical "
        "investor. The target user may understand returns but may not naturally check volatility, "
        "drawdown, turnover, concentration, or sentiment quality before choosing a fund. The app "
        "therefore starts with the strongest evidence, supports comparison and detailed fund "
        "inspection, and then lets the investor test an allocation after fees. Sentiment and method "
        "evidence remain close to the decision so the user can see why a result should be treated "
        "cautiously.",
    ),
    ("Heading 1", "Funds and Backtest Design"),
    ("Heading 2", "Funds and Optimisation Methods"),
    (
        "Normal",
        "BetaVest tests twelve funds across three groups of permitted assets: equity-only, "
        "crypto-only, and combined equity-plus-crypto. Each group is tested using four methods: equal "
        "weight, minimum variance, maximum Sharpe, and risk parity. This satisfies the required "
        "combined-fund minimum and adds equity-only and crypto-only sleeves, which helps the "
        "report explain whether crypto is best used alone or as part of a diversified fund.",
    ),
    (
        "Normal",
        "Each asset-family and method pair is treated as one fund because this is what the "
        "investor would choose in the app. For example, Combined Risk Parity is a different "
        "fund from Equity Risk Parity because it is allowed to invest in different assets and has a different risk "
        "profile and holding set. This structure also makes the fact sheets clearer: each "
        "fund has its own growth path, drawdown history, performance metrics, current holdings, "
        "and suggested investor use.",
    ),
    (
        "Normal",
        "The four methods represent different investment philosophies. Equal weight is the "
        "transparent benchmark because it does not estimate expected returns or covariances. "
        "Minimum variance targets lower realised volatility and drawdown. Maximum Sharpe targets "
        "return per unit of estimated risk, but it is more sensitive to noisy historical means. "
        "Risk parity uses scipy.optimize.minimize to solve a long-only equal-risk-contribution "
        "problem, with weights constrained to be non-negative and to sum to one. The risk parity "
        "fund is therefore included as a more stable optimisation method, not only as another "
        "return-maximising rule.",
    ),
    (
        "Normal",
        "Using several methods is useful because each one can fail in a different way. Equal weight "
        "can over-allocate to volatile assets because it ignores risk. Minimum variance can become "
        "too conservative if it avoids assets with useful expected returns. Maximum Sharpe can chase "
        "assets that looked strong in the estimation window but do not continue to outperform. Risk "
        "parity can still be exposed to common market shocks, but it is less dependent on unstable "
        "return forecasts. Comparing these methods side by side gives a stronger answer than choosing "
        "one optimiser in advance.",
    ),
    (
        "Normal",
        "The permitted-asset split is also part of the design. Equity-only funds show what could be "
        "achieved using the 50-stock dataset. Crypto-only funds show the standalone risk "
        "of digital assets. Combined funds test whether crypto improves a broader portfolio once the "
        "optimiser is allowed to manage risk across both families. This matters because the investor "
        "decision is not simply whether crypto has high returns; it is whether crypto improves the "
        "portfolio after accounting for volatility, drawdown, and diversification.",
    ),
    ("Heading 2", "Data Alignment and Investable Timing"),
    (
        "Normal",
        "The backtest uses the cleaned Part A return panel. Equity returns are calculated from "
        "adjusted-close prices on the US equity trading calendar. Crypto returns are calculated "
        "on their native daily calendar and then aligned to equity trading dates for combined "
        "funds. This avoids turning weekend crypto moves into artificial weekday equity trades "
        "and ensures that the combined funds are evaluated on a common investable calendar.",
    ),
    (
        "Normal",
        "The backtest is walk-forward and out of sample. At each month-end rebalance, the model "
        "estimates weights using the previous 252 trading days. The first live backtest date is "
        "therefore after the initial estimation window rather than the first date in the dataset. "
        "Weights are formed using information available before the rebalance date, then the "
        "following returns are used to calculate realised performance. This timing rule is the "
        "main protection against look-ahead bias.",
    ),
    (
        "Normal",
        "Sharpe ratios use a zero risk-free rate, which is stated as a simplifying assumption. "
        "Performance is reported after a 10-basis-point transaction cost on turnover. Including "
        "turnover matters because some optimisers improve gross return by trading more frequently. "
        "A retail-facing fund product should not rank funds only by pre-cost performance if the "
        "implementation cost changes the final investor experience.",
    ),
    (
        "Normal",
        "The rebalance design deliberately favours interpretability over overfitting. Monthly "
        "rebalancing is frequent enough for the funds to respond to changing risk and sentiment, "
        "but not so frequent that the project becomes a daily trading strategy. The 252-day lookback "
        "window approximates one trading year, giving the optimiser enough observations to estimate "
        "risk while still adapting to recent market conditions. These choices are simple, but they "
        "are defensible and easy for a marker or investor to audit.",
    ),
    (
        "Normal",
        "The main backtest risk is that historical estimates can look precise even when they are "
        "not. To reduce this risk, the report does not rely on one metric. Annualised return shows "
        "growth potential, volatility shows day-to-day risk, Sharpe ratio combines return and risk, "
        "maximum drawdown captures the worst investor pain point, turnover shows implementation "
        "intensity, and transaction cost shows whether trading reduces the practical result. A fund "
        "only earns a strong recommendation when the full set of measures is consistent.",
    ),
    ("Heading 1", "Out-of-Sample Results and Fund Fact Sheets"),
    ("Heading 2", "Main Results"),
    (
        "Normal",
        "The strongest out-of-sample fund is Combined Risk Parity. It records an annualised "
        "return of 13.65%, annualised volatility of 16.04%, a Sharpe ratio of 0.851, a maximum "
        "drawdown of -19.79%, and a total return of 46.57%. This is the main fund-design answer: "
        "the best core BetaVest fund is not the most aggressive crypto portfolio, but a combined "
        "fund that controls how much each asset contributes to risk.",
    ),
    (
        "Normal",
        "Combined Maximum Sharpe has the highest annualised return among the main combined funds "
        "at 16.18% and the highest total return at 56.52%. However, it also has higher annualised "
        "volatility of 22.39%, a deeper maximum drawdown of -26.01%, and total transaction cost "
        "of 1.80%. It is therefore better described as an aggressive growth sleeve than as the "
        "default core fund. Combined Equal Weight also performs well, with 14.94% annualised "
        "return and a 0.703 Sharpe ratio, but the -28.75% drawdown is less attractive for a "
        "young investor who may not understand large peak-to-trough losses.",
    ),
    (
        "Normal",
        "The equity-only funds provide useful benchmarks. Equity Equal Weight delivers a 12.61% "
        "annualised return and a 0.780 Sharpe ratio, making it a strong transparent benchmark. "
        "Equity Risk Parity produces a lower annualised return of 9.80%, but it has lower "
        "volatility of 14.58% and a smaller maximum drawdown of -18.68%. Equity Minimum Variance "
        "is more defensive, with 12.77% volatility and a -15.55% maximum drawdown, but its "
        "annualised return falls to 4.25%. These funds show the return-risk trade-off clearly: "
        "lower volatility is useful, but it can come at a large opportunity cost.",
    ),
    (
        "Normal",
        "The crypto-only funds are not suitable as core BetaVest funds in this sample. Crypto "
        "Minimum Variance has the highest annualised return in the full results table at 22.61%, "
        "but annualised volatility is 64.95% and maximum drawdown is -75.13%. Crypto Equal Weight "
        "has a weak Sharpe ratio of 0.040 and a maximum drawdown of -84.85%. Crypto Maximum "
        "Sharpe performs worst, losing 49.53% over the live period. These results support a "
        "clear product message: crypto can be used as a satellite exposure or a controlled part "
        "of a combined fund, but crypto-only products should be labelled speculative.",
    ),
    (
        "Normal",
        "The 2022 drawdown is especially important for interpretation. A fund can look attractive "
        "in a full-sample table if strong recovery periods offset one severe loss period. However, "
        "a real investor experiences the path, not only the final average. The drawdown view shows "
        "whether the investor would have needed to tolerate large losses before the fund recovered. "
        "This is why Combined Risk Parity is preferred even though Combined Maximum Sharpe and Crypto "
        "Minimum Variance produce higher raw return numbers in some summaries.",
    ),
    (
        "Normal",
        "The results also show why BetaVest should avoid a purely mechanical ranking table. If the "
        "app ranked funds only by annualised return, a user could select Crypto Minimum Variance "
        "without recognising that it lost more than half its value in the 2022 period. If the app "
        "ranked funds only by minimum drawdown, the user might select a defensive fund that gives "
        "up too much upside. The better decision rule is to combine performance, risk, role, and "
        "investor suitability.",
    ),
    ("Heading 2", "Fact-Sheet Interpretation"),
    (
        "Normal",
        "The fact sheets translate these results into investor-facing decisions. Combined Risk "
        "Parity should be presented as the core growth candidate because it balances return, "
        "volatility, drawdown, and cost better than the other funds. Equity Equal Weight should "
        "be presented as the simple equity benchmark. Combined Maximum Sharpe should be presented "
        "as an aggressive growth sleeve because it raises return but also increases volatility, "
        "drawdown, and turnover. Equity Minimum Variance should be presented as a defensive "
        "stabiliser because it reduces risk but gives up too much return to be the main offer.",
    ),
    (
        "Normal",
        "This use-case framing is more useful than a single performance ranking. A young retail "
        "investor may incorrectly select the highest-return fund without understanding drawdown "
        "or volatility. The fact sheet design avoids this by showing growth of one dollar, "
        "annualised return, annualised volatility, Sharpe ratio, maximum drawdown, total return, "
        "current holdings, asset-class exposure, sector exposure, holding concentration, turnover, "
        "and transaction cost. These are the key checks needed before treating any fund as a "
        "core allocation.",
    ),
    (
        "Normal",
        "The current holdings section is included because a fund label can hide concentration. "
        "Two funds with similar Sharpe ratios may expose the investor to very different assets or "
        "sectors. Showing the latest holdings, asset-class split, sector exposure, top-five weight, "
        "and top-ten weight makes the fact sheet closer to a real fund document. It helps the user "
        "ask whether the fund is diversified, whether crypto exposure is material, and whether one "
        "sector is driving too much of the result.",
    ),
    (
        "Normal",
        "The fact sheets also connect the report to the Streamlit app. In the report, the fact-sheet "
        "evidence supports the written recommendation. In the app, the same evidence becomes an "
        "interactive decision page where the user can choose a fund from a dropdown and immediately "
        "see its risk profile. This link between written analysis and dashboard workflow is important "
        "because the deliverable is not only a report; it is a report, codebase, and investor-facing "
        "application.",
    ),
    (
        "Normal",
        "The out-of-sample evidence therefore supports a balanced product recommendation. "
        "BetaVest should lead with Combined Risk Parity, retain Equity Equal Weight as the "
        "simple benchmark, offer Combined Maximum Sharpe for users who explicitly accept higher "
        "risk, and treat crypto-only funds as satellite exposure only. This recommendation follows "
        "from risk-adjusted evidence rather than from the fund with the most eye-catching return.",
    ),
    ("Heading 1", "Sentiment Index"),
    ("Heading 2", "Construction"),
    (
        "Normal",
        "The sentiment index converts the Part A headline dataset into a standalone equity-sector "
        "news signal. The input data contains headline text, ticker labels, sector labels, and "
        "aligned trading dates. The model is deliberately transparent: it is designed as a finance "
        "headline-tone index rather than a black-box article-level classifier. Sentiment is applied "
        "only to equities because the project news data covers the permitted equities, not the crypto "
        "assets.",
    ),
    (
        "Normal",
        "The scoring model uses a finance-lexicon baseline. Positive words include terms such as "
        "beat, buy, growth, upgrade, gain, strong, profit, outperform, rally, and record. Negative "
        "words include miss, sell, downgrade, risk, loss, weak, decline, lawsuit, cut, warning, "
        "debt, and probe. Before scoring, headlines are lowercased and punctuation is removed so "
        "finance terms can be matched consistently. Numbers and ticker symbols are not treated as "
        "sentiment words because the model is a transparent lexicon signal rather than a full "
        "language classifier.",
    ),
    (
        "Normal",
        "For each ticker-day, the score is positive word counts minus negative word counts, scaled "
        "by the number of sentiment-bearing terms. Sector sentiment is then calculated by averaging "
        "ticker-day scores across the five stocks in each sector. Ticker-days with no headlines are "
        "treated as neutral. This is a conservative choice because it avoids allowing one heavily "
        "covered firm to dominate the sector index when the other firms have no news. The signal is "
        "lagged by one trading day before investment use, so day t sentiment can only affect a day "
        "t+1 decision or later.",
    ),
    (
        "Normal",
        "The neutral treatment for missing headlines is a judgement call, but it is appropriate for "
        "this project. Dropping missing ticker-days would make the index depend only on firms that "
        "received coverage, which could exaggerate event-driven news. Carrying sentiment forward "
        "would assume yesterday's tone remains valid even when there is no new information. Neutral "
        "scoring is conservative: it says that no observed headline should not automatically become "
        "positive or negative evidence. This choice makes the sentiment index less dramatic but more "
        "suitable for a controlled portfolio tilt.",
    ),
    (
        "Normal",
        "The one-day lag is the most important implementation detail. Headlines must be available "
        "before they can be used in a trading decision. A same-day signal would risk using information "
        "that was published after the portfolio decision or after market close. By lagging the score, "
        "the project treats sentiment as information known before the next rebalance. This keeps the "
        "text signal consistent with the out-of-sample backtest logic used for the return-based funds.",
    ),
    ("Heading 2", "Findings and Investment Use"),
    (
        "Normal",
        "The sentiment summary shows that news tone differs across sectors. Consumer has the highest "
        "average sentiment at 0.521, followed by Technology at 0.485 and Healthcare at 0.396. "
        "Materials has the lowest average sentiment at 0.164. Coverage also differs: Technology has "
        "26,632 aligned headlines and Consumer has 25,877, while Materials has 5,393 and Real Estate "
        "has 5,339. These differences mean the index should be read as both a sentiment measure and "
        "a coverage-aware signal.",
    ),
    (
        "Normal",
        "The useful investment signal is the lagged sector sentiment index. It can tilt an equity "
        "portfolio toward sectors with stronger recent news tone, flag sectors where the news tone "
        "is weakening, and help explain why the equity sleeve changes over time. It should not be "
        "read as a guarantee that positive news causes positive returns. Sentiment is noisy, and a "
        "small lexicon can miss negation and context. The index is therefore most credible when it "
        "supports an existing risk-controlled portfolio process.",
    ),
    (
        "Normal",
        "This design answers the sentiment requirement in the brief. The project applies a sentiment "
        "model to the assembled headlines, builds a sector-level sentiment index, equal-weights "
        "ticker-day scores within sectors, treats missing ticker-days neutrally, and lags the signal "
        "before trading. These choices make the signal reproducible and reduce the risk of overstating "
        "what the text model can know at the time of investment.",
    ),
    (
        "Normal",
        "The sector sentiment figure is useful because it shows persistence and disagreement across "
        "sectors rather than only a final average. A sector with high average sentiment but unstable "
        "weekly movement may be less useful for allocation than a sector with a smoother signal. The "
        "selected-sector table in the app helps with this interpretation by updating to the sectors "
        "chosen by the user. This makes the sentiment section more than a static exhibit: it becomes "
        "a way to inspect where the news signal is strongest, weakest, and most reliable.",
    ),
    (
        "Normal",
        "The app also adapts the Week 9 fear-greed idea into a BetaVest sentiment gauge. Instead "
        "of presenting the raw sector sentiment score alone, the latest 21-day average market "
        "sentiment is converted into a 0-100 percentile against BetaVest history. A low score "
        "therefore means current news sentiment is unusually weak for this dataset, while a high "
        "score means it is unusually strong. The gauge is placed in the Sentiment tab beside the "
        "sector index and tilt evidence so it is read as an investor mood check, not as a direct "
        "buy-or-sell rule or as the full CNN fear-greed methodology.",
    ),
    ("Heading 1", "Extensions and Innovations"),
    ("Heading 2", "Sentiment Fusion"),
    (
        "Normal",
        "The main extension is a sentiment-fusion fund that adds the lagged sector sentiment signal "
        "to Equity Risk Parity. The base fund first estimates long-only risk-parity weights using "
        "the previous 252 trading days. At each monthly rebalance, each equity then receives a "
        "sector sentiment multiplier based on one-day-lagged sentiment. The multiplier is capped "
        "between 0.70 and 1.30 and weights are renormalised to sum to 100%. This keeps the fund "
        "long-only and prevents the sentiment signal from overwhelming the optimiser.",
    ),
    (
        "Normal",
        "The full-sample fusion result is positive but modest. Equity Risk Parity has annualised "
        "return of 9.80%, volatility of 14.58%, a Sharpe ratio of 0.672, a maximum drawdown of "
        "-18.68%, and total transaction cost of 0.23%. Equity Sentiment Tilt improves annualised "
        "return to 10.08%, Sharpe to 0.690, and maximum drawdown to -18.27%, while total transaction "
        "cost rises to 0.46%. The full-sample Sharpe lift is 0.018, so the evidence supports "
        "sentiment as an incremental tilt rather than a complete investment model.",
    ),
    (
        "Normal",
        "The tilt-strength grid is separated from the final test. The 2021-2022 discovery period "
        "selects a finance-lexicon tilt of 0.15, where Sharpe rises from 0.572 for Equity Risk Parity "
        "to 0.576 for the tilt. That choice is frozen before evaluating 2023. In the untouched "
        "holdout, annualised return rises from 11.53% to 12.23% and Sharpe rises from 0.944 to 1.003, "
        "a lift of 0.059 after transaction costs. This separation reduces the risk that the reported "
        "result comes from selecting a parameter that happened to fit the full sample.",
    ),
    (
        "Normal",
        "A key point is that the sentiment fusion is intentionally constrained. If the model allowed "
        "large sentiment bets, the portfolio could look innovative but become fragile. The cap on "
        "the multiplier, the long-only constraint, and the monthly rebalance rule keep the signal "
        "inside an investable fund design. This is why the result is reported as a Sharpe lift rather "
        "than as a claim that sentiment can forecast every stock return. The value is incremental, "
        "measured, and linked to a baseline fund.",
    ),
    ("Heading 2", "Sentiment Model Comparison"),
    (
        "Normal",
        "The project compares the finance lexicon with standard VADER before treating sentiment as "
        "reliable. The finance lexicon has mean ticker-day sentiment of 0.436, compared with 0.279 "
        "for VADER. Their score correlation is only 0.337, and they agree on positive, neutral, or "
        "negative direction for 58.8% of ticker-days. The finance lexicon classifies 41.3% of "
        "ticker-days as neutral, while VADER classifies 24.1% as neutral. Model choice therefore "
        "materially changes the measured news signal.",
    ),
    (
        "Normal",
        "Both models produce a positive 2023 holdout result after their tilt strengths are selected "
        "on 2021-2022. The finance lexicon raises holdout Sharpe by 0.059, while VADER raises it by "
        "0.037. The finance model remains the selected BetaVest signal, but the result is reported "
        "beside the VADER benchmark and a reproducible sample of the largest model disagreements. "
        "The disagreement sample is a review tool rather than labelled ground truth; a production "
        "model would still require independent human coding of headline tone.",
    ),
    ("Heading 2", "Robustness and Regime Analysis"),
    (
        "Normal",
        "The second extension is a robustness test across market periods. The backtest is split "
        "into the full sample, the 2021 recovery, the 2022 drawdown, and the 2023 recovery. Combined "
        "Risk Parity remains a strong core candidate because it performs well in the recovery years "
        "while limiting downside in 2022 better than more speculative alternatives. Its Sharpe ratio "
        "is 2.501 in 2021, -0.354 in 2022, and 1.224 in 2023. Crypto Minimum Variance has strong "
        "upside in favourable years, but its 2022 return of -57.58% and maximum drawdown of -62.91% "
        "confirm that crypto-only exposure fails the stability test for core use.",
    ),
    (
        "Normal",
        "The third extension is sentiment regime analysis. Days are grouped into negative, neutral, "
        "and positive regimes using terciles of the one-day-lagged sector sentiment signal. Equity "
        "Sentiment Tilt improves Equity Risk Parity Sharpe in the negative regime from 0.516 to "
        "0.527 and in the positive regime from 0.200 to 0.261. It underperforms in the neutral "
        "regime, where Sharpe falls from 1.445 to 1.406. This result is balanced: the sentiment signal helps "
        "in some regimes, but it is not a universal return predictor.",
    ),
    (
        "Normal",
        "Together, these extensions address the higher-band requirement in the brief. The project "
        "adds equity-only and crypto-only funds, implements an extra optimisation method through "
        "long-only risk parity, adds transaction costs and turnover, builds a sentiment-fusion fund, "
        "selects tilt strength on a discovery period, tests an untouched holdout, compares two text "
        "models, reports robustness by market period, and analyses performance across sentiment "
        "regimes. The result is not just a larger set of funds; it is a more honest "
        "evaluation of when each fund should and should not be used.",
    ),
    (
        "Normal",
        "The robustness work is also important because it prevents the innovation section from "
        "becoming only a feature list. The project does not simply add a new signal and stop after "
        "one full-sample result. It asks whether the signal survives different market conditions "
        "and whether the best fund remains credible during stress. That makes the extension more "
        "evidence-based and closer to the standard expected in a fund research note.",
    ),
    ("Heading 1", "App and Investor Journey"),
    ("Heading 2", "Dashboard Structure"),
    (
        "Normal",
        "The Streamlit app converts the modelling outputs into an investor-facing dashboard. It "
        "does not recompute the raw data pipeline at runtime. Instead, it reads precomputed result "
        "files from the project results folder, which keeps the deployed app fast, reproducible, "
        "and easier to audit. The app includes an Overview tab, Compare tab, Fact Sheet tab, "
        "Allocation tab, Sentiment tab, and Method tab.",
    ),
    (
        "Normal",
        "The Overview tab explains what BetaVest does, reports the main findings, summarises the "
        "number of funds and assets covered, shows a featured fund, points users to the sentiment "
        "analytics, and explains the risk-tier labels. The Compare tab then lets the user select funds "
        "to overlay and switch between growth of one dollar, drawdown, and rolling Sharpe. A global "
        "date-range control in the sidebar lets the same sample window flow through the dashboard, "
        "and the app uses consistent fund colours so the same fund is recognisable across tabs.",
    ),
    (
        "Normal",
        "The Fact Sheet tab supports deeper inspection of one fund. It shows return, volatility, "
        "Sharpe ratio, drawdown, total return, turnover, cost, worst drawdown period, benchmark "
        "comparison, holdings, asset-class exposure, sector exposure, and holding concentration. "
        "The Allocation tab then lets the user combine selected funds using sliders and exact "
        "percentage inputs. A management-fee control compares gross and net growth of the selected "
        "mix, making the cost of the product visible. The Sentiment tab exposes the fear-greed "
        "gauge, sector sentiment index, model comparison, discovery grid, 2023 holdout, and "
        "sentiment regime analysis. The Method tab explains "
        "the formulas and caveats so the app is not a black box.",
    ),
    (
        "Normal",
        "Several app choices are designed to improve usability rather than add unnecessary decoration. "
        "The sidebar keeps global settings such as chart horizon, date range, data status, and theme "
        "in one place. Fund, fact-sheet, and sector selectors sit on the relevant pages so they are "
        "visible when the user needs them. A consistent colour is assigned to each fund across charts, "
        "allocation views, and comparison tables. The dark theme is included for readability, while "
        "download buttons let the user export the underlying evidence.",
    ),
    ("Heading 2", "Investor Journey"),
    (
        "Normal",
        "The intended investor journey has four steps. First, the user starts on Overview to see "
        "what BetaVest offers and what the headline result is. Second, the user moves to Compare "
        "to test whether the featured fund is still attractive under different performance views "
        "and date ranges. Third, the user opens the Fact Sheet to inspect holdings and risks before "
        "choosing the fund. Fourth, the user tests a simple allocation and checks the Sentiment and "
        "Method tabs to understand whether the recommendation is supported by both market data and "
        "headline evidence.",
    ),
    (
        "Normal",
        "This journey matters because the target user is a young retail investor, not a portfolio "
        "research team. The app therefore uses risk labels, suggested uses, readable cards, a "
        "light/dark theme option, global controls, downloadable tables, and concise subtitles under "
        "each tab. These features help the user understand the analysis without reading the code. "
        "They also make the app closer to a product: it guides fund discovery, evidence checking, "
        "allocation building, fee evaluation, and method review in one place.",
    ),
    (
        "Normal",
        "The investor journey is deliberately ordered. Overview gives orientation, Compare helps the "
        "user decide which funds deserve attention, Fact Sheet checks whether one fund is acceptable, "
        "Allocation turns fund choices into a portfolio mix, Sentiment explains the text signal behind "
        "the equity tilt, and Method documents the assumptions. This sequence reduces cognitive load "
        "because the user is not asked to understand optimisation formulas before seeing the practical "
        "fund evidence.",
    ),
    (
        "Normal",
        "The app still avoids pretending to be personalised financial advice. It presents evidence "
        "from historical simulations and makes the caveats visible. This is why drawdown, volatility, "
        "turnover, transaction cost, and robustness are shown beside returns. The app's strongest "
        "design feature is that it slows down the decision process: before selecting a fund, the user "
        "can see what the fund owns, how it behaved in the 2022 drawdown, and whether sentiment "
        "evidence adds anything meaningful.",
    ),
    ("Heading 1", "Critical Reflection and Recommendations"),
    ("Heading 2", "Reflection"),
    (
        "Normal",
        "The first limitation is estimation error. Maximum-Sharpe and mean-variance methods use "
        "historical expected returns and covariances, which are noisy in a short sample. The 2020-2023 "
        "period contains pandemic stress, high-growth technology repricing, inflation pressure, and "
        "large crypto cycles. These events make optimiser inputs unstable. Risk parity reduces this "
        "problem because it focuses on risk contribution rather than estimated mean returns, but it "
        "does not remove estimation error entirely.",
    ),
    (
        "Normal",
        "The second limitation is sentiment measurement. A transparent finance lexicon is useful "
        "because the scoring rule is easy to inspect, but it can miss negation, sarcasm, multi-word "
        "phrases, and context. For example, the same word can have different meanings across sectors "
        "or headline formats. The sentiment result should therefore be treated as evidence that a "
        "simple signal can add incremental value, not as proof that the current lexicon is production "
        "ready. The low cross-model correlation and 58.8% direction agreement confirm that the "
        "measured signal depends on the scoring model.",
    ),
    (
        "Normal",
        "The third limitation is deployment realism. The app is reproducible because it reads "
        "precomputed result files, but the final submission still needs a public Streamlit URL, a "
        "public repository link, and a clean AI workflow pack. In a real product, users would also "
        "expect more recent data, live monitoring, clearer disclaimers, and stronger stress testing. "
        "The current project is a strong coursework prototype, not a regulated investment product.",
    ),
    (
        "Normal",
        "There are also limitations in the sample period. The live backtest covers 2020-2023 after "
        "the initial estimation window, which includes major market stress and recovery episodes but "
        "is still short for judging long-term fund behaviour. A different interest-rate cycle, a "
        "different crypto regime, or a different equity sector leadership pattern could change the "
        "rankings. This is why the conclusion should be read as evidence from this sample, not as a "
        "permanent ranking of investment strategies.",
    ),
    ("Heading 2", "Three Concrete Recommendations"),
    (
        "Normal",
        "Recommendation 1 is to strengthen the optimiser inputs. The next version should add "
        "shrinkage covariance estimates, maximum holding caps, and concentration diagnostics. This "
        "would reduce the chance that one noisy return window drives unrealistic weights. It would "
        "also make the fact sheets more credible because users could see whether a fund is diversified "
        "or dependent on a small number of assets.",
    ),
    (
        "Normal",
        "Recommendation 2 is to complete independent headline validation. The project now compares "
        "the finance lexicon with VADER and produces a targeted disagreement sample, but model "
        "agreement is not accuracy. A human-coded sample should identify false positives, false "
        "negatives, and sector-specific language before the lexicon is expanded. The revised signal "
        "should then repeat the same 2021-2022 discovery and 2023 holdout test.",
    ),
    (
        "Normal",
        "Recommendation 3 is to complete the product and hand-in deployment checklist. The final "
        "version should include the live Streamlit app URL, public GitHub repository link, standard "
        "app entrypoint, updated Word fields, exported PDF, and AI workflow pack. The app should also "
        "retain the gross-versus-net fee view and include clear disclaimers that results are historical simulations and not personalised "
        "financial advice.",
    ),
    (
        "Normal",
        "These recommendations are concrete because each one can be implemented and evaluated. "
        "Shrinkage covariance and holding caps can be tested by rerunning the same walk-forward "
        "backtest. Sentiment validation can be measured by comparing labelled headline samples with "
        "model scores and by retesting the tilt after lexicon changes. Deployment improvements can "
        "be checked directly through the app URL, repository structure, downloadable files, and user "
        "journey. This makes the reflection actionable rather than only descriptive.",
    ),
    (
        "Normal",
        "Overall, BetaVest answers the Part B brief by building systematic funds, testing them out "
        "of sample, creating a sector sentiment index, fusing sentiment with an equity optimiser, "
        "adding robustness checks, and presenting the evidence in a usable app. The strongest result "
        "is Combined Risk Parity as the core fund candidate. The most useful innovation is not that "
        "sentiment wins every test, but that the project shows where sentiment helps, where it does "
        "not, and how it can be controlled inside a disciplined investment process.",
    ),
]


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_before(reference, text: str, style: str):
    paragraph = reference.insert_paragraph_before(text)
    paragraph.style = style
    paragraph.paragraph_format.space_after = Pt(5)
    if style == "Heading 1":
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
    elif style == "Heading 2":
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def count_narrative_words() -> int:
    return sum(len(re.findall(r"\b[\w.-]+\b", text)) for style, text in NARRATIVE if style == "Normal")


def compact_last_appendix_table(doc: Document) -> None:
    """Reduce the final long appendix table so it does not leave an orphan row."""
    if not doc.tables:
        return
    table = doc.tables[-1]
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 0.85
                for run in paragraph.runs:
                    run.font.size = Pt(6.5)


def main() -> None:
    doc = Document(REPORT)

    for paragraph in list(doc.paragraphs[:5]):
        if paragraph.style.name == "Title" and not paragraph.text.strip():
            delete_paragraph(paragraph)

    reference_index = next(
        i
        for i, paragraph in enumerate(doc.paragraphs)
        if paragraph.style.name == "Heading 1" and paragraph.text.strip() == "References"
    )
    intro_index = next(
        i
        for i, paragraph in enumerate(doc.paragraphs)
        if paragraph.style.name == "Heading 1" and paragraph.text.strip() == "Introduction"
    )
    reference = doc.paragraphs[reference_index]

    for paragraph in list(doc.paragraphs[intro_index:reference_index]):
        delete_paragraph(paragraph)

    for style, text in NARRATIVE:
        insert_before(reference, text, style)

    word_count = count_narrative_words()
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Word Count:"):
            paragraph.text = f"Word Count: approximately {word_count:,} words excluding appendix and references"
            break

    compact_last_appendix_table(doc)
    doc.save(REPORT)
    print(f"Updated {REPORT} with approximately {word_count:,} narrative words.")


if __name__ == "__main__":
    main()
