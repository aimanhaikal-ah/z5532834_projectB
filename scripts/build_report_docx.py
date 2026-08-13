"""Build report/report.docx from the Markdown draft and generated exhibits."""
from __future__ import annotations

import pathlib
import re

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


ROOT = pathlib.Path(__file__).resolve().parent.parent
SOURCE = ROOT / "report" / "PART_B_DRAFT.md"
OUTPUT = ROOT / "report" / "report.docx"
FIGURES = ROOT / "results" / "figures"
TABLES = ROOT / "results" / "tables"

EXHIBITS = {
    "Table 1.": ("table", TABLES / "performance_metrics_report.csv"),
    "Table 2.": ("table", TABLES / "sector_sentiment_summary.csv"),
    "Table 3.": ("table", TABLES / "fusion_comparison_report.csv"),
    "Table 4.": ("table", TABLES / "sentiment_tilt_grid_report.csv"),
    "Table 5.": ("table", TABLES / "robustness_by_period_appendix.csv"),
    "Table 6.": ("table", TABLES / "sentiment_regime_analysis_appendix.csv"),
    "Table 7.": ("table", TABLES / "sentiment_holdout_validation_appendix.csv"),
    "Table 8.": ("table", TABLES / "sentiment_model_comparison_report.csv"),
    "Figure 1.": ("figure", FIGURES / "fund_growth_of_1.png"),
    "Figure 2.": ("figure", FIGURES / "selected_fund_drawdown.png"),
    "Figure 3.": ("figure", FIGURES / "selected_fund_weights_over_time.png"),
    "Figure 4.": ("figure", FIGURES / "fund_sharpe_ratios.png"),
    "Figure 5.": ("figure", FIGURES / "sector_sentiment_index.png"),
    "Figure 6.": ("figure", FIGURES / "fusion_base_vs_sentiment.png"),
}


def strip_md(text: str) -> str:
    """Remove the small amount of Markdown markup used in the draft."""
    return text.replace("**", "").replace("`", "")


def fmt_value(value, column: str) -> str:
    """Format report table values to 2-3 significant digits."""
    if pd.isna(value):
        return ""
    if column in {
        "Ann. return (%)",
        "Ann. volatility (%)",
        "Max drawdown (%)",
        "Total return (%)",
        "Total cost (%)",
        "Average daily return (%)",
        "Positive return days (%)",
        "Neutral ticker-days (%)",
        "Positive ticker-days (%)",
        "Negative ticker-days (%)",
        "Direction agreement (%)",
        "Holdout base return (%)",
        "Holdout tilt return (%)",
        "Holdout base drawdown (%)",
        "Holdout tilt drawdown (%)",
        "Holdout tilt cost (%)",
    }:
        return f"{float(value):.2f}"
    if column in {
        "Sharpe",
        "Average sentiment",
        "Avg tickers with headlines",
        "Avg daily turnover",
        "Tilt strength",
        "Average lagged sentiment",
        "Selected tilt",
        "Discovery base Sharpe",
        "Discovery tilt Sharpe",
        "Discovery Sharpe lift",
        "Holdout base Sharpe",
        "Holdout tilt Sharpe",
        "Holdout Sharpe lift",
        "Mean sentiment",
        "Standard deviation",
        "Cross-model correlation",
    }:
        return f"{float(value):.3f}"
    if column in {"Total headlines", "Observations"}:
        return f"{int(value):,}"
    return str(value)


def add_dataframe(doc: Document, csv_path: pathlib.Path) -> None:
    """Insert a compact Word table from a CSV."""
    df = pd.read_csv(csv_path)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, column in enumerate(df.columns):
        cell = table.rows[0].cells[i]
        cell.text = column
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(7.5)

    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, column in enumerate(df.columns):
            cells[i].text = fmt_value(row[column], column)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)

    doc.add_paragraph()


def add_figure(doc: Document, image_path: pathlib.Path) -> None:
    """Insert a generated PNG figure."""
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(6.4))
    doc.add_paragraph()


def apply_styles(doc: Document) -> None:
    """Apply simple report styles."""
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
    styles["Caption"].font.name = "Arial"
    styles["Caption"].font.size = Pt(9)
    styles["Caption"].font.italic = True


def add_exhibit_if_needed(doc: Document, caption: str) -> None:
    """Insert the matching table or figure after a caption."""
    for prefix, (kind, path) in EXHIBITS.items():
        if caption.startswith(prefix):
            if kind == "table":
                add_dataframe(doc, path)
            else:
                add_figure(doc, path)
            return


def add_references_section(doc: Document) -> None:
    """Add a clean references section to the Word output."""
    doc.add_page_break()
    doc.add_paragraph("References", style="Heading 1")
    paragraph = doc.add_paragraph(
        "The empirical results in this report are generated from the submitted "
        "BetaVest Part B project code, result tables, and figures. Key project "
        "artifacts include scripts/run_part_b.py, streamlit_app.py, and the "
        "generated files under results/."
    )
    paragraph.paragraph_format.space_after = Pt(5)


def add_appendix_section(doc: Document, captions: list[str]) -> None:
    """Append generated tables and figures after the report body."""
    if not captions:
        return
    table_captions = [caption for caption in captions if caption.startswith("Table")]
    figure_captions = [caption for caption in captions if caption.startswith("Figure")]
    doc.add_page_break()
    doc.add_paragraph("Appendix", style="Heading 1")
    if table_captions:
        doc.add_paragraph("Appendix A: Tables", style="Heading 2")
    for caption in table_captions:
        paragraph = doc.add_paragraph(caption)
        paragraph.style = "Caption"
        add_exhibit_if_needed(doc, caption)
    if figure_captions:
        doc.add_page_break()
        doc.add_paragraph("Appendix B: Figures", style="Heading 2")
    for caption in figure_captions:
        paragraph = doc.add_paragraph(caption)
        paragraph.style = "Caption"
        add_exhibit_if_needed(doc, caption)


def build() -> pathlib.Path:
    """Create the Word report and apply the final expanded narrative."""
    doc = Document()
    apply_styles(doc)
    appendix_captions: list[str] = []
    skip_markdown_working_notes = False

    for line in SOURCE.read_text().splitlines():
        text = line.strip()
        if not text:
            continue
        if text in {"## Required Exhibits Checklist", "## Human Edit Notes"}:
            skip_markdown_working_notes = True
            continue
        if skip_markdown_working_notes:
            continue

        if text.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.style = "Title"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(strip_md(text[2:]))
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(18)
        elif text.startswith("## "):
            doc.add_paragraph(strip_md(text[3:]), style="Heading 1")
        elif text.startswith("### "):
            doc.add_paragraph(strip_md(text[4:]), style="Heading 2")
        elif text.startswith("- [ ] "):
            doc.add_paragraph(strip_md(text[6:]), style="List Bullet")
        elif text.startswith("- "):
            doc.add_paragraph(strip_md(text[2:]), style="List Bullet")
        elif re.match(r"^\*\*(Table|Figure) [0-9]+\.", text):
            appendix_captions.append(strip_md(text))
        elif text.startswith("**") and text.endswith("**"):
            paragraph = doc.add_paragraph(strip_md(text))
            paragraph.style = "Caption"
        else:
            paragraph = doc.add_paragraph(strip_md(text))
            paragraph.paragraph_format.space_after = Pt(5)

    add_references_section(doc)
    add_appendix_section(doc, appendix_captions)
    doc.save(OUTPUT)

    # The Markdown file remains a planning draft. Apply the fuller Word-first
    # narrative last so a normal report rebuild cannot overwrite it.
    from rewrite_report_narrative import main as rewrite_narrative

    rewrite_narrative()
    return OUTPUT


if __name__ == "__main__":
    print(build())
